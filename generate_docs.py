import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfgen import canvas

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)
        print(f"REPORTLAB PDF PAGE COUNT: {num_pages}")


# ----------------------------------------------------
# 1. GENERATE DOCX FILE (SINGLE COLUMN LAYOUT)
# ----------------------------------------------------

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_docx(filename):
    doc = docx.Document()
    
    # 0.5 in margins all around for compact single-page layout
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Styles
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # --- HEADER TABLE ---
    h_table = doc.add_table(rows=1, cols=2)
    h_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_table.autofit = False
    
    c_left, c_right = h_table.rows[0].cells[0], h_table.rows[0].cells[1]
    c_left.width = Inches(4.5)
    c_right.width = Inches(2.6)
    
    p_title = c_left.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("SERVICE AGREEMENT")
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p_sub = c_left.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(0)
    r_sub = p_sub.add_run("DOCUMENTARY VIDEO EDITING & PRODUCTION")
    r_sub.font.size = Pt(9)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    p_meta = c_right.paragraphs[0]
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_meta.paragraph_format.space_after = Pt(0)
    r1 = p_meta.add_run("Effective Date: ")
    r1.font.bold = True
    r1.font.size = Pt(9)
    p_meta.add_run("Immediately\n")
    r2 = p_meta.add_run("Valid Until: ")
    r2.font.bold = True
    r2.font.size = Pt(9)
    p_meta.add_run("31 January 2027")

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(4)
    p_div.paragraph_format.space_after = Pt(8)
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="1" w:color="0F172A"/></w:pBdr>')
    p_div._p.get_or_add_pPr().append(pBdr)

    # --- PARTIES BANNER ---
    p_table = doc.add_table(rows=1, cols=2)
    p_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_table.autofit = False
    
    cell_ed, cell_cl = p_table.rows[0].cells[0], p_table.rows[0].cells[1]
    cell_ed.width, cell_cl.width = Inches(3.55), Inches(3.55)
    
    for cell, role, name, sub in [
        (cell_ed, "EDITOR", "Sufiyan", "SufiVisuals"),
        (cell_cl, "CLIENT", "Arshmeet singh", "Client Partner")
    ]:
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        set_cell_border(cell,
                        top={'val':'single', 'sz':4, 'color':'CBD5E1'},
                        bottom={'val':'single', 'sz':4, 'color':'CBD5E1'},
                        left={'val':'single', 'sz':4, 'color':'CBD5E1'},
                        right={'val':'single', 'sz':4, 'color':'CBD5E1'})
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r_lbl = p.add_run(f"{role}: ")
        r_lbl.font.size = Pt(8.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        
        r_nm = p.add_run(name)
        r_nm.font.size = Pt(11)
        r_nm.font.bold = True
        r_nm.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        if sub:
            r_sb = p.add_run(f" ({sub})")
            r_sb.font.size = Pt(9)
            r_sb.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Helper for Section Headings
    def add_section_title(num, title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        
        r_num = p.add_run(f"{num}. ")
        r_num.font.size = Pt(11)
        r_num.font.bold = True
        r_num.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        r_title = p.add_run(title_text.upper())
        r_title.font.size = Pt(10.5)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)

    # --- SECTION 1: SCOPE OF WORK & DELIVERABLES ---
    add_section_title("01", "Scope of Work & Deliverables")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Scope of Work: ")
    r.font.bold = True
    p.add_run("Documentary-style YouTube video editing. Approximately 1 video per week (depending on content schedule), with an average video duration of 10–20 minutes.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Editing Scope (when required): ")
    r.font.bold = True
    p.add_run("Storytelling & pacing, motion graphics, maps & route animations, B-roll integration, AI visuals, text animations, sound design, and basic color correction.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Final Deliverables: ")
    r.font.bold = True
    p.add_run("Final exported video in 4K or 1080p resolution as required. ")
    r_files = p.add_run("Project files are not included unless agreed separately.")
    r_files.font.italic = True

    # --- SECTION 2: PRICING, REVISIONS & SCOPE CHANGES ---
    add_section_title("02", "Pricing, Revisions & Scope Changes")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Pricing: ")
    r.font.bold = True
    r_pr = p.add_run("₹9,000 per completed video.")
    r_pr.font.bold = True
    r_pr.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Revisions & Turnaround: ")
    r.font.bold = True
    p.add_run("Up to ")
    r_2 = p.add_run("2 revision rounds")
    r_2.font.bold = True
    p.add_run(" per video. Additional revisions or scope alterations may incur extra charges. Turnaround time is mutually agreed for each video based on complexity.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Scope Changes: ")
    r.font.bold = True
    p.add_run("If a future video requires significantly more work than the current documentary editing style (for example: extensive VFX, 3D animation, advanced motion design, or additional research beyond the agreed scope), pricing will be discussed separately before work begins.")

    # --- SECTION 3: CLIENT RESPONSIBILITIES & PAYMENT ---
    add_section_title("03", "Client Responsibilities & Payment Terms")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Client Responsibilities: ")
    r.font.bold = True
    p.add_run("Provide script, voice-over, raw footage, branding assets (if any), and reference material before editing begins. Provide consolidated feedback for revision rounds.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Payment Terms: ")
    r.font.bold = True
    p.add_run("Payment to be made after delivery of each completed video unless otherwise agreed. Any applicable taxes or transaction charges are the client's responsibility.")

    # --- SECTION 4: TERMINATION, CONFIDENTIALITY & OWNERSHIP ---
    add_section_title("04", "Termination, Confidentiality & Ownership")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Termination: ")
    r.font.bold = True
    p.add_run("Either party may end the agreement with ")
    r_7 = p.add_run("7 days' written notice")
    r_7.font.bold = True
    p.add_run(". Completed work must be paid for before termination.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Confidentiality: ")
    r.font.bold = True
    p.add_run("Both parties agree to keep confidential files, business information, project materials, and unpublished content private.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Ownership & Portfolio Showcase: ")
    r.font.bold = True
    p.add_run("Ownership of the final edited video transfers to the client after full payment is received. The editor retains the right to showcase completed work in their portfolio unless the client requests confidentiality in writing.")

    # --- SECTION 5: SIGNATURES ---
    add_section_title("05", "Signatures & Execution")
    
    p_sig_intro = doc.add_paragraph()
    p_sig_intro.paragraph_format.space_after = Pt(4)
    p_sig_intro.add_run("By signing below, both parties confirm and accept all terms of this agreement.")

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    s1, s2 = sig_table.rows[0].cells[0], sig_table.rows[0].cells[1]
    s1.width, s2.width = Inches(3.55), Inches(3.55)
    
    for cell, role, name in [
        (s1, "EDITOR", "Sufiyan (SufiVisuals)"),
        (s2, "CLIENT", "Arshmeet singh")
    ]:
        set_cell_background(cell, "FAFAFA")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        set_cell_border(cell,
                        top={'val':'single', 'sz':4, 'color':'E2E8F0'},
                        bottom={'val':'single', 'sz':4, 'color':'E2E8F0'},
                        left={'val':'single', 'sz':4, 'color':'E2E8F0'},
                        right={'val':'single', 'sz':4, 'color':'E2E8F0'})
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r_role = p.add_run(role)
        r_role.font.size = Pt(8)
        r_role.font.bold = True
        r_role.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        
        p_name = cell.add_paragraph()
        p_name.paragraph_format.space_after = Pt(12)
        r_nm = p_name.add_run(name)
        r_nm.font.size = Pt(10)
        r_nm.font.bold = True
        r_nm.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        p_line = cell.add_paragraph()
        p_line.paragraph_format.space_after = Pt(2)
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="94A3B8"/></w:pBdr>')
        p_line._p.get_or_add_pPr().append(pBdr)
        
        p_meta = cell.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(0)
        r_sig = p_meta.add_run("Authorized Signature")
        r_sig.font.size = Pt(7.5)
        r_sig.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        
        r_dt = p_meta.add_run("                                    Date: ____________")
        r_dt.font.size = Pt(7.5)
        r_dt.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.save(filename)
    print(f"DOCX saved to {filename}")


# ----------------------------------------------------
# 2. GENERATE PDF FILE (SINGLE COLUMN, 1 PAGE GUARANTEED)
# ----------------------------------------------------

def create_pdf(filename):
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        leftMargin=38,
        rightMargin=38,
        topMargin=30,
        bottomMargin=30
    )

    styles = getSampleStyleSheet()
    
    primary = colors.HexColor('#0F172A')
    secondary = colors.HexColor('#475569')
    muted = colors.HexColor('#64748B')
    border_col = colors.HexColor('#CBD5E1')
    bg_light = colors.HexColor('#F8FAFC')
    
    style_title = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=17,
        leading=19,
        textColor=primary
    )
    
    style_subtitle = ParagraphStyle(
        'DocSub',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=10,
        textColor=muted
    )

    style_meta = ParagraphStyle(
        'DocMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11,
        alignment=2,
        textColor=secondary
    )

    style_sec_title = ParagraphStyle(
        'SecTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=12,
        textColor=primary
    )

    style_body = ParagraphStyle(
        'SingleBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.2,
        leading=12,
        textColor=colors.HexColor('#1E293B')
    )

    elements = []

    # --- HEADER TABLE ---
    h_data = [
        [
            Paragraph("<b>SERVICE AGREEMENT</b><br/><font color='#64748B' size=8>DOCUMENTARY VIDEO EDITING & PRODUCTION</font>", style_title),
            Paragraph("<b>Effective Date:</b> Immediately<br/><b>Valid Until:</b> 31 January 2027", style_meta)
        ]
    ]
    t_header = Table(h_data, colWidths=[330, 189])
    t_header.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ('LINEBELOW', (0,0), (-1,-1), 1.5, primary),
    ]))
    elements.append(t_header)
    elements.append(Spacer(1, 6))

    # --- PARTIES BANNER ---
    party_left = Paragraph("<font size=7 color='#64748B'><b>EDITOR:</b></font> <font size=10 color='#0F172A'><b>Sufiyan</b></font> <font size=8.5 color='#475569'>(SufiVisuals)</font>", style_body)
    party_right = Paragraph("<font size=7 color='#64748B'><b>CLIENT:</b></font> <font size=10 color='#0F172A'><b>Arshmeet singh</b></font> <font size=8.5 color='#475569'>(Client Partner)</font>", style_body)
    
    t_parties = Table([[party_left, party_right]], colWidths=[259, 260])
    t_parties.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), bg_light),
        ('BOX', (0,0), (0,0), 0.5, border_col),
        ('BOX', (1,0), (1,0), 0.5, border_col),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    elements.append(t_parties)
    elements.append(Spacer(1, 6))

    def make_sec_header(num, title):
        h_table = Table([[
            Paragraph(f"<b>{num}. {title.upper()}</b>", style_sec_title)
        ]], colWidths=[519])
        h_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 2),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ('LINEBELOW', (0,0), (-1,-1), 0.5, border_col),
        ]))
        return h_table

    # --- SECTION 1 ---
    elements.append(make_sec_header("01", "Scope of Work & Deliverables"))
    elements.append(Spacer(1, 3))
    elements.append(Paragraph("• <b>Scope of Work:</b> Documentary-style YouTube video editing. Approximately 1 video per week (depending on content schedule), with an average video duration of 10–20 minutes.", style_body))
    elements.append(Spacer(1, 2))
    elements.append(Paragraph("• <b>Editing Scope (when required):</b> Storytelling & pacing, motion graphics, maps & route animations, B-roll integration, AI visuals, text animations, sound design, and basic color correction.", style_body))
    elements.append(Spacer(1, 2))
    elements.append(Paragraph("• <b>Final Deliverables:</b> Final exported video in 4K or 1080p resolution as required. <i>Project files are not included unless agreed separately.</i>", style_body))
    elements.append(Spacer(1, 6))

    # --- SECTION 2 ---
    elements.append(make_sec_header("02", "Pricing, Revisions & Scope Changes"))
    elements.append(Spacer(1, 3))
    elements.append(Paragraph("• <b>Agreed Pricing:</b> <b><font color='#0F172A'>₹9,000 per completed video.</font></b>", style_body))
    elements.append(Spacer(1, 2))
    elements.append(Paragraph("• <b>Revisions & Turnaround:</b> Up to <b>2 revision rounds</b> per video. Additional revisions or scope alterations may incur extra charges. Turnaround time is mutually agreed for each video based on complexity.", style_body))
    elements.append(Spacer(1, 2))
    elements.append(Paragraph("• <b>Scope Changes:</b> If a future video requires significantly more work than the current documentary editing style (for example: extensive VFX, 3D animation, advanced motion design, or additional research beyond the agreed scope), pricing will be discussed separately before work begins.", style_body))
    elements.append(Spacer(1, 6))

    # --- SECTION 3 ---
    elements.append(make_sec_header("03", "Client Responsibilities & Payment Terms"))
    elements.append(Spacer(1, 3))
    elements.append(Paragraph("• <b>Client Responsibilities:</b> Provide script, voice-over, raw footage, branding assets (if any), and reference material before editing begins. Provide consolidated feedback for revision rounds.", style_body))
    elements.append(Spacer(1, 2))
    elements.append(Paragraph("• <b>Payment Terms:</b> Payment to be made after delivery of each completed video unless otherwise agreed. Any applicable taxes or transaction charges are the client's responsibility.", style_body))
    elements.append(Spacer(1, 6))

    # --- SECTION 4 ---
    elements.append(make_sec_header("04", "Termination, Confidentiality & Ownership"))
    elements.append(Spacer(1, 3))
    elements.append(Paragraph("• <b>Termination:</b> Either party may end the agreement with <b>7 days' written notice</b>. Completed work must be paid for before termination.", style_body))
    elements.append(Spacer(1, 2))
    elements.append(Paragraph("• <b>Confidentiality:</b> Both parties agree to keep confidential files, business information, project materials, and unpublished content private.", style_body))
    elements.append(Spacer(1, 2))
    elements.append(Paragraph("• <b>Ownership & Portfolio:</b> Ownership of the final edited video transfers to the client after full payment is received. The editor retains the right to showcase completed work in their portfolio unless the client requests confidentiality in writing.", style_body))
    elements.append(Spacer(1, 8))

    # --- SECTION 5: SIGNATURES ---
    elements.append(make_sec_header("05", "Signatures & Execution"))
    elements.append(Spacer(1, 4))
    elements.append(Paragraph("By signing below, both parties confirm and accept all terms of this agreement.", ParagraphStyle('SubSig', parent=style_body, fontSize=8.5, textColor=secondary)))
    elements.append(Spacer(1, 4))

    sig_editor = [
        Paragraph("<font size=7 color='#64748B'><b>EDITOR</b></font><br/><b><font size=9.5 color='#0F172A'>Sufiyan (SufiVisuals)</font></b>", style_body),
        Spacer(1, 14),
        Paragraph("<font color='#94A3B8'>________________________________________</font><br/><font size=7 color='#64748B'>Authorized Signature &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; Date</font>", style_body)
    ]

    sig_client = [
        Paragraph("<font size=7 color='#64748B'><b>CLIENT</b></font><br/><b><font size=9.5 color='#0F172A'>Arshmeet singh</font></b>", style_body),
        Spacer(1, 14),
        Paragraph("<font color='#94A3B8'>________________________________________</font><br/><font size=7 color='#64748B'>Authorized Signature &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; Date</font>", style_body)
    ]

    t_sig = Table([[sig_editor, sig_client]], colWidths=[259, 260])
    t_sig.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#FAFAFA')),
        ('BOX', (0,0), (0,0), 0.5, border_col),
        ('BOX', (1,0), (1,0), 0.5, border_col),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    elements.append(t_sig)

    # Build PDF
    doc.build(elements, canvasmaker=NumberedCanvas)
    print(f"PDF saved to {filename}")

if __name__ == '__main__':
    docx_file = r"d:\frame alternative\Service_Agreement_SufiVisuals_ArshmeetSingh.docx"
    pdf_file = r"d:\frame alternative\Service_Agreement_SufiVisuals_ArshmeetSingh.pdf"
    
    create_docx(docx_file)
    create_pdf(pdf_file)
